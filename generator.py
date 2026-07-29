from fpdf import FPDF
from docx import Document
from models import MeetingNotes, Segment


def _to_latin1(text: str) -> str:
    """Coerce text to Latin-1 so fpdf2's core Helvetica font can render it.

    Meeting transcripts occasionally contain CJK / emoji / other chars that the
    core Type-1 fonts cannot encode. Rather than bundling a Unicode TTF (adds
    several MB to the .app), replace unrepresentable code points with '?'.
    """
    if text is None:
        return ""
    return str(text).encode("latin-1", errors="replace").decode("latin-1")


class _MeetingPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 16)
        self.cell(0, 10, "Meeting Notes", new_x="LMARGIN", new_y="NEXT", align="C")
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    def _write_section(self, title: str, items: list[str]):
        self.set_font("Helvetica", "B", 12)
        self.cell(0, 10, _to_latin1(title), new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 10)
        for item in items:
            self.cell(6, 6, "-")
            self.multi_cell(0, 6, _to_latin1(item), new_x="LMARGIN", new_y="NEXT")
        self.ln(4)


def generate_pdf(notes: MeetingNotes) -> bytes:
    pdf = _MeetingPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 10, "Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, _to_latin1(notes.summary))
    pdf.ln(6)

    pdf._write_section("Key Discussion Points", notes.key_points)
    pdf._write_section("Action Items", notes.action_items)
    pdf._write_section("Decisions Made", notes.decisions)
    pdf._write_section("Topics Covered", notes.topics)

    return bytes(pdf.output())


def generate_text(notes: MeetingNotes) -> str:
    lines = [
        "MEETING NOTES",
        "=" * 60,
        "",
        "SUMMARY",
        "-" * 20,
        notes.summary,
        "",
        "KEY DISCUSSION POINTS",
        "-" * 25,
    ]
    for p in notes.key_points:
        lines.append(f"  \u2022 {p}")

    lines.extend(["", "ACTION ITEMS", "-" * 20])
    for a in notes.action_items:
        lines.append(f"  \u2022 {a}")

    lines.extend(["", "DECISIONS MADE", "-" * 20])
    for d in notes.decisions:
        lines.append(f"  \u2022 {d}")

    lines.extend(["", "TOPICS COVERED", "-" * 20])
    for t in notes.topics:
        lines.append(f"  \u2022 {t}")

    lines.extend(["", "FULL TRANSCRIPT", "-" * 20, notes.transcript])

    return "\n".join(lines)


def _fmt_timestamp(seconds: float) -> str:
    total = int(seconds)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _transcript_lines(notes: MeetingNotes) -> list[tuple[str, str]]:
    """Return (timestamp, text) pairs, falling back to the plain transcript if no segments."""
    if notes.segments:
        return [(_fmt_timestamp(s.start), s.text) for s in notes.segments]
    if notes.transcript:
        return [("", notes.transcript)]
    return []


class _TranscriptPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 16)
        self.cell(0, 10, "Transcript", new_x="LMARGIN", new_y="NEXT", align="C")
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def generate_transcript_pdf(notes: MeetingNotes) -> bytes:
    pdf = _TranscriptPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    for ts, text in _transcript_lines(notes):
        if ts:
            pdf.set_font("Helvetica", "B", 9)
            pdf.write(5, _to_latin1(f"[{ts}] "))
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, _to_latin1(text), new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, _to_latin1(text), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    return bytes(pdf.output())


def generate_transcript_docx(notes: MeetingNotes) -> bytes:
    import io

    doc = Document()
    doc.add_heading("Transcript", level=1)

    for ts, text in _transcript_lines(notes):
        p = doc.add_paragraph()
        if ts:
            run = p.add_run(f"[{ts}] ")
            run.bold = True
        p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
